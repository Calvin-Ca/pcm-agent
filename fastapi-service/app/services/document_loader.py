"""
文档加载器模块

支持加载PDF、Word、Markdown和TXT文档
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)


class DocumentChunk:
    """文档块"""
    
    def __init__(
        self,
        content: str,
        metadata: Dict[str, Any],
        chunk_id: Optional[str] = None
    ):
        self.content = content
        self.metadata = metadata
        self.chunk_id = chunk_id or f"chunk_{hash(content)}"
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "metadata": self.metadata
        }


class DocumentLoader:
    """文档加载器"""
    
    def __init__(self, max_workers: int = 4):
        self.max_workers = max_workers
        self.executor = ThreadPoolExecutor(max_workers=max_workers)
    
    async def load_document(self, file_path: str) -> List[DocumentChunk]:
        """
        加载单个文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            文档块列表
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文档不存在: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        # 根据文件类型选择加载器
        if file_ext == '.pdf':
            return await self._load_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await self._load_docx(file_path)
        elif file_ext in ['.md', '.markdown']:
            return await self._load_markdown(file_path)
        elif file_ext == '.txt':
            return await self._load_text(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")
    
    async def load_documents(self, file_paths: List[str]) -> List[DocumentChunk]:
        """
        批量加载文档
        
        Args:
            file_paths: 文档路径列表
            
        Returns:
            所有文档块列表
        """
        tasks = [self.load_document(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        all_chunks = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"加载文档失败 {file_paths[i]}: {result}")
                continue
            all_chunks.extend(result)
        
        return all_chunks
    
    async def _load_pdf(self, file_path: str) -> List[DocumentChunk]:
        """加载PDF文档"""
        if PyPDF2 is None:
            raise ImportError("需要安装PyPDF2: pip install PyPDF2")
        
        def _extract_pdf_text():
            chunks = []
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text.strip():
                                metadata = {
                                    "source": file_path,
                                    "page": page_num + 1,
                                    "type": "pdf",
                                    "total_pages": len(pdf_reader.pages)
                                }
                                chunks.append(DocumentChunk(
                                    content=text.strip(),
                                    metadata=metadata
                                ))
                        except Exception as e:
                            logger.warning(f"提取PDF第{page_num + 1}页失败: {e}")
                            continue
                            
            except Exception as e:
                logger.error(f"读取PDF文件失败: {e}")
                raise
            
            return chunks
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract_pdf_text)
    
    async def _load_docx(self, file_path: str) -> List[DocumentChunk]:
        """加载Word文档"""
        if DocxDocument is None:
            raise ImportError("需要安装python-docx: pip install python-docx")
        
        def _extract_docx_text():
            chunks = []
            try:
                doc = DocxDocument(file_path)
                
                # 提取段落
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                if paragraphs:
                    content = '\n\n'.join(paragraphs)
                    metadata = {
                        "source": file_path,
                        "type": "docx",
                        "paragraphs": len(paragraphs)
                    }
                    chunks.append(DocumentChunk(
                        content=content,
                        metadata=metadata
                    ))
                
                # 提取表格
                for table_idx, table in enumerate(doc.tables):
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        table_text.append(' | '.join(row_text))
                    
                    if table_text:
                        content = '\n'.join(table_text)
                        metadata = {
                            "source": file_path,
                            "type": "docx_table",
                            "table_index": table_idx
                        }
                        chunks.append(DocumentChunk(
                            content=content,
                            metadata=metadata
                        ))
                        
            except Exception as e:
                logger.error(f"读取Word文档失败: {e}")
                raise
            
            return chunks
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract_docx_text)
    
    async def _load_markdown(self, file_path: str) -> List[DocumentChunk]:
        """加载Markdown文档"""
        def _extract_markdown_text():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                
                if content.strip():
                    metadata = {
                        "source": file_path,
                        "type": "markdown",
                        "size": len(content)
                    }
                    return [DocumentChunk(
                        content=content.strip(),
                        metadata=metadata
                    )]
                return []
                
            except Exception as e:
                logger.error(f"读取Markdown文件失败: {e}")
                raise
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract_markdown_text)
    
    async def _load_text(self, file_path: str) -> List[DocumentChunk]:
        """加载文本文档"""
        def _extract_text():
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                
                if content.strip():
                    metadata = {
                        "source": file_path,
                        "type": "text",
                        "size": len(content)
                    }
                    return [DocumentChunk(
                        content=content.strip(),
                        metadata=metadata
                    )]
                return []
                
            except Exception as e:
                logger.error(f"读取文本文件失败: {e}")
                raise
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, _extract_text)
    
    def close(self):
        """关闭线程池"""
        self.executor.shutdown(wait=True)


class ChunkSplitter:
    """文档分块器"""
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        separators: Optional[List[str]] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ['\n\n', '\n', '。', '！', '？', '.', '!', '?']
    
    def split_chunks(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """
        分割文档块
        
        Args:
            chunks: 原始文档块列表
            
        Returns:
            分割后的文档块列表
        """
        split_chunks = []
        
        for chunk in chunks:
            if len(chunk.content) <= self.chunk_size:
                split_chunks.append(chunk)
            else:
                # 分割长文档
                sub_chunks = self._split_text(chunk.content, chunk.metadata)
                split_chunks.extend(sub_chunks)
        
        return split_chunks
    
    def _split_text(self, text: str, base_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """分割文本"""
        chunks = []
        
        # 尝试按分隔符分割
        for separator in self.separators:
            if separator in text:
                parts = text.split(separator)
                current_chunk = ""
                
                for part in parts:
                    # 如果添加这部分不会超过限制
                    if len(current_chunk) + len(part) + len(separator) <= self.chunk_size:
                        if current_chunk:
                            current_chunk += separator + part
                        else:
                            current_chunk = part
                    else:
                        # 保存当前块
                        if current_chunk.strip():
                            metadata = base_metadata.copy()
                            metadata["chunk_index"] = len(chunks)
                            chunks.append(DocumentChunk(
                                content=current_chunk.strip(),
                                metadata=metadata
                            ))
                        
                        # 开始新块（包含重叠）
                        if self.chunk_overlap > 0 and current_chunk:
                            overlap_text = current_chunk[-self.chunk_overlap:]
                            current_chunk = overlap_text + separator + part
                        else:
                            current_chunk = part
                
                # 添加最后一块
                if current_chunk.strip():
                    metadata = base_metadata.copy()
                    metadata["chunk_index"] = len(chunks)
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        metadata=metadata
                    ))
                
                return chunks
        
        # 如果没有找到合适的分隔符，按字符数强制分割
        return self._split_by_chars(text, base_metadata)
    
    def _split_by_chars(self, text: str, base_metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """按字符数强制分割"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            if chunk_text.strip():
                metadata = base_metadata.copy()
                metadata["chunk_index"] = len(chunks)
                metadata["forced_split"] = True
                chunks.append(DocumentChunk(
                    content=chunk_text.strip(),
                    metadata=metadata
                ))
            
            # 移动到下一个位置（考虑重叠）
            start = end - self.chunk_overlap if self.chunk_overlap > 0 else end
        
        return chunks