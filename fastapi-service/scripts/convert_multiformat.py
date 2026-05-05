"""
多格式转换脚本：将 knowledge-base-v2/ 中的 .md 转换为 .docx 和 .pdf，并生成 .csv。

分布目标：.md ≈ 70 / .docx 18 / .pdf 10 / .csv 2 = 100 篇

用法:
    cd fastapi-service
    python -m scripts.convert_multiformat
"""
import random
import re
import csv
from pathlib import Path
from collections import defaultdict

# 确保可复现
random.seed(42)

KB_ROOT = Path("E:/huan/工时管理系统/trunk/1 源代码/1.0 系统代码/ai-service/knowledge-base-v2")


def extract_frontmatter(md_text: str) -> dict:
    """提取 frontmatter 字段"""
    fm = {}
    m = re.match(r"^---\s*\n(.*?)\n---\s*\n", md_text, re.DOTALL)
    if m:
        for line in m.group(1).split("\n"):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if ":" in line:
                k, _, v = line.partition(":")
                fm[k.strip()] = v.strip()
    return fm


def md_to_docx(md_path: Path, out_path: Path):
    """将 .md 转为 .docx，保留 frontmatter 作为文档属性"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = md_path.read_text(encoding="utf-8")
    fm = extract_frontmatter(text)
    body = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, count=1, flags=re.DOTALL).strip()

    doc = Document()

    # 标题
    title = fm.get("title", md_path.stem)
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Frontmatter 信息表
    if fm:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "属性"
        hdr_cells[1].text = "值"
        for k, v in fm.items():
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[1].text = v
        doc.add_paragraph()

    # 正文（简易 markdown 解析）
    for line in body.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
            doc.add_paragraph(line[3:], style="List Number")
        elif line.startswith("```"):
            continue  # 跳过代码块标记
        else:
            doc.add_paragraph(line)

    doc.save(str(out_path))
    print(f"  [DOCX] {md_path.name} -> {out_path.name}")


def md_to_pdf(md_path: Path, out_path: Path):
    """将 .md 转为 .pdf，使用 reportlab"""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 尝试注册中文字体
    font_name = "Helvetica"
    try:
        # 尝试常见 Windows 中文字体
        font_paths = [
            "C:/Windows/Fonts/simhei.ttf",   # 黑体
            "C:/Windows/Fonts/simsun.ttc",   # 宋体
            "C:/Windows/Fonts/msyh.ttc",     # 微软雅黑
        ]
        for fp in font_paths:
            if Path(fp).exists():
                font_name = "ChineseFont"
                pdfmetrics.registerFont(TTFont(font_name, fp))
                break
    except Exception as e:
        print(f"  [WARN] 中文字体注册失败: {e}, 使用默认字体")
        font_name = "Helvetica"

    text = md_path.read_text(encoding="utf-8")
    fm = extract_frontmatter(text)
    body = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, count=1, flags=re.DOTALL).strip()

    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ChineseBody",
        fontName=font_name,
        fontSize=10,
        leading=16,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="ChineseHeading2",
        fontName=font_name,
        fontSize=14,
        leading=20,
        spaceAfter=10,
        textColor=colors.HexColor("#1a1a1a"),
    ))
    styles.add(ParagraphStyle(
        name="ChineseHeading3",
        fontName=font_name,
        fontSize=12,
        leading=16,
        spaceAfter=8,
        textColor=colors.HexColor("#333333"),
    ))
    styles.add(ParagraphStyle(
        name="ChineseTitle",
        fontName=font_name,
        fontSize=18,
        leading=24,
        alignment=1,  # center
        spaceAfter=20,
        textColor=colors.HexColor("#000000"),
    ))

    story = []

    # 标题
    title = fm.get("title", md_path.stem)
    story.append(Paragraph(title, styles["ChineseTitle"]))
    story.append(Spacer(1, 0.5*cm))

    # Frontmatter 表格
    if fm:
        data = [["属性", "值"]]
        for k, v in fm.items():
            data.append([k, v])
        t = Table(data, colWidths=[4*cm, 12*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("FONTNAME", (0, 0), (-1, 0), font_name),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 1), (-1, -1), font_name),
            ("FONTSIZE", (0, 1), (-1, -1), 9),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

    # 正文
    for line in body.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("## "):
            story.append(Paragraph(line[3:], styles["ChineseHeading2"]))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles["ChineseHeading3"]))
        elif line.startswith("- "):
            story.append(Paragraph("• " + line[2:], styles["ChineseBody"]))
        elif line.startswith("```"):
            continue
        else:
            # 转义 HTML 特殊字符
            line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(line, styles["ChineseBody"]))

    doc.build(story)
    print(f"  [PDF]  {md_path.name} -> {out_path.name}")


def select_files_for_conversion(all_md_files: list, docx_count: int, pdf_count: int) -> tuple:
    """
    选择要转换的文件，确保每个主题域至少保留一些 .md，
    且转换的文件分布到各个主题域。
    """
    # 按主题域分组
    by_domain = defaultdict(list)
    for f in all_md_files:
        # 路径如: knowledge-base-v2/01-工时管理/policy/xxx.md
        parts = f.relative_to(KB_ROOT).parts
        domain = parts[0] if parts else "unknown"
        by_domain[domain].append(f)

    domains = sorted(by_domain.keys())
    print(f"主题域: {domains}")
    for d in domains:
        print(f"  {d}: {len(by_domain[d])} 篇")

    # 每个主题域至少选 2 篇用于转换（1 docx + 1 pdf 或 2 docx 等）
    docx_selected = []
    pdf_selected = []

    # 先每个主题域至少分配 1 篇 docx 和 1 篇 pdf（如果够的话）
    for domain in domains:
        files = by_domain[domain]
        if len(files) >= 2:
            picks = random.sample(files, min(3, len(files)))
            docx_selected.append(picks[0])
            pdf_selected.append(picks[1])
        elif len(files) == 1:
            docx_selected.append(files[0])

    # 去重
    docx_selected = list(dict.fromkeys(docx_selected))
    pdf_selected = list(dict.fromkeys(pdf_selected))

    # 从剩余文件中补充到目标数量
    remaining_for_docx = [f for f in all_md_files if f not in docx_selected]
    remaining_for_pdf = [f for f in all_md_files if f not in pdf_selected and f not in docx_selected]

    needed_docx = docx_count - len(docx_selected)
    if needed_docx > 0 and remaining_for_docx:
        extra = random.sample(remaining_for_docx, min(needed_docx, len(remaining_for_docx)))
        docx_selected.extend(extra)

    needed_pdf = pdf_count - len(pdf_selected)
    if needed_pdf > 0 and remaining_for_pdf:
        extra = random.sample(remaining_for_pdf, min(needed_pdf, len(remaining_for_pdf)))
        pdf_selected.extend(extra)

    return docx_selected[:docx_count], pdf_selected[:pdf_count]


def generate_csv_files():
    """生成 2 个 CSV 文件"""
    # 1. 2026年假期日历
    holiday_dir = KB_ROOT / "02-假期与加班" / "data"
    holiday_dir.mkdir(parents=True, exist_ok=True)

    holidays = [
        ("2026-01-01", "元旦", "1", "无", ""),
        ("2026-02-17", "春节", "7", "2026-02-15,2026-02-21", "除夕至初六"),
        ("2026-04-05", "清明节", "3", "无", "含周末连休"),
        ("2026-05-01", "劳动节", "5", "2026-04-26", "含周末调休"),
        ("2026-06-19", "端午节", "3", "无", "含周末连休"),
        ("2026-09-25", "中秋节", "3", "无", "含周末连休"),
        ("2026-10-01", "国庆节", "7", "2026-09-20", "10月1日至7日"),
    ]

    csv_path = holiday_dir / "2026年假期日历.csv"
    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["日期", "假期类型", "放假天数", "补班日期", "备注"])
        writer.writerows(holidays)
    print(f"  [CSV] 生成 {csv_path}")

    # 2. 项目审批节点表
    project_dir = KB_ROOT / "06-项目管理流程" / "data"
    project_dir.mkdir(parents=True, exist_ok=True)

    approvals = [
        ("小型项目", "部门经理", "项目经理", "部门经理", "3"),
        ("中型项目", "事业部总监", "部门经理", "事业部总监", "5"),
        ("大型项目", "公司副总", "事业部总监", "公司副总", "10"),
        ("战略项目", "CEO", "公司副总", "CEO", "15"),
    ]

    csv_path2 = project_dir / "项目审批节点表.csv"
    with open(csv_path2, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["项目类型", "立项审批人", "变更审批人", "验收审批人", "SLA(工作日)"])
        writer.writerows(approvals)
    print(f"  [CSV] 生成 {csv_path2}")


def main():
    print("=" * 60)
    print("多格式转换开始")
    print("=" * 60)

    # 1. 收集所有 .md 文件
    all_md_files = list(KB_ROOT.rglob("*.md"))
    print(f"\n找到 {len(all_md_files)} 篇 .md 文件")

    # 2. 选择要转换的文件
    docx_files, pdf_files = select_files_for_conversion(all_md_files, 18, 10)
    print(f"\n选中转换:")
    print(f"  .docx: {len(docx_files)} 篇")
    print(f"  .pdf:  {len(pdf_files)} 篇")

    # 3. 转换 .docx
    print("\n--- 转换为 .docx ---")
    for md_path in docx_files:
        # 文件名编码 category 和 genre
        rel = md_path.relative_to(KB_ROOT)
        parts = rel.parts
        domain = parts[0] if len(parts) > 0 else ""
        genre = parts[1] if len(parts) > 1 else ""
        stem = md_path.stem
        # 新文件名: 原文件名__genre_category.docx
        new_name = f"{stem}__{genre}_{domain}.docx"
        out_path = md_path.with_name(new_name).with_suffix(".docx")
        md_to_docx(md_path, out_path)
        # 删除原 .md
        md_path.unlink()
        print(f"  [DEL] 删除原 {md_path.name}")

    # 4. 转换 .pdf
    print("\n--- 转换为 .pdf ---")
    for md_path in pdf_files:
        rel = md_path.relative_to(KB_ROOT)
        parts = rel.parts
        domain = parts[0] if len(parts) > 0 else ""
        genre = parts[1] if len(parts) > 1 else ""
        stem = md_path.stem
        new_name = f"{stem}__{genre}_{domain}.pdf"
        out_path = md_path.with_name(new_name).with_suffix(".pdf")
        md_to_pdf(md_path, out_path)
        # 删除原 .md
        md_path.unlink()
        print(f"  [DEL] 删除原 {md_path.name}")

    # 5. 生成 CSV
    print("\n--- 生成 .csv ---")
    generate_csv_files()

    # 6. 统计最终分布
    print("\n" + "=" * 60)
    print("转换完成，最终分布统计:")
    print("=" * 60)

    md_count = len(list(KB_ROOT.rglob("*.md")))
    docx_count_actual = len(list(KB_ROOT.rglob("*.docx")))
    pdf_count_actual = len(list(KB_ROOT.rglob("*.pdf")))
    csv_count_actual = len(list(KB_ROOT.rglob("*.csv")))
    total = md_count + docx_count_actual + pdf_count_actual + csv_count_actual

    print(f"  .md:   {md_count}")
    print(f"  .docx: {docx_count_actual}")
    print(f"  .pdf:  {pdf_count_actual}")
    print(f"  .csv:  {csv_count_actual}")
    print(f"  总计:  {total}")

    # 按主题域统计
    print("\n按主题域分布:")
    for domain_dir in sorted(KB_ROOT.iterdir()):
        if not domain_dir.is_dir():
            continue
        d_md = len(list(domain_dir.rglob("*.md")))
        d_docx = len(list(domain_dir.rglob("*.docx")))
        d_pdf = len(list(domain_dir.rglob("*.pdf")))
        d_csv = len(list(domain_dir.rglob("*.csv")))
        d_total = d_md + d_docx + d_pdf + d_csv
        print(f"  {domain_dir.name}: md={d_md} docx={d_docx} pdf={d_pdf} csv={d_csv} (total={d_total})")

    # 按格式统计
    print("\n按格式统计:")
    by_ext = defaultdict(int)
    for f in KB_ROOT.rglob("*"):
        if f.is_file() and f.suffix in (".md", ".docx", ".pdf", ".csv"):
            by_ext[f.suffix] += 1
    for ext, count in sorted(by_ext.items()):
        print(f"  {ext}: {count}")

    return total


if __name__ == "__main__":
    main()
